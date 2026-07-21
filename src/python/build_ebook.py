"""Build artifacts/reports/Leatt_Growth_OS_Data_Story.docx — the polished
print/PDF-style companion to index.html.

Design rule for this rebuild: only use screenshots verified as genuine
(real browser chrome, real portal captures) — leatt_about_source,
leatt_moto_boots_collection, leatt_lifestyle_eyewear_collection,
azure_portal_home_initial/costs. Codex's powerbi_*.png / competitor_positioning.png
etc. are matplotlib mockups mislabeled as if captured from the Power BI
service (one, powerbi_executive_overview.png, even shows numbers
inconsistent with the real gold_monthly_kpis.csv) — deliberately excluded
rather than propagated. Charts are the ones this rebuild generated from the
real CSVs (src/python/build_charts.py), verified against source data first.

KPIs are computed from leatt_intelligent_category/channel/monthly/province.csv
— the correctly-scaled gold aggregates matching the real 2,000,000-row fact
table (verified: sums to R3,994,247,412.56, matching leatt_reconciliation.csv's
sales_incl_vat_zar total exactly, and matching every figure already cited in
leatt_kpi_rationale_catalog.csv / leatt_decision_signal_rules.csv /
leatt_next_best_actions.csv). The old gold_*_kpis.csv files were a stale,
~400x-smaller demo sample that produced numbers inconsistent with those
strategic files — not used here.
"""

import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

BASE = os.path.join(os.path.dirname(__file__), "..", "..")
DATA = os.path.join(BASE, "artifacts", "data_samples")
IMG = os.path.join(BASE, "artifacts", "evidence_images")
CHARTS = os.path.join(IMG, "charts")
OUT = os.path.join(BASE, "artifacts", "reports", "Leatt_Growth_OS_Data_Story.docx")

RED = RGBColor(0xD7, 0x19, 0x20)
INK = RGBColor(0x14, 0x14, 0x14)
GREY = RGBColor(0x6E, 0x6E, 0x6E)
GOLD = RGBColor(0xB8, 0x86, 0x0B)


def money(v):
    if abs(v) >= 1e9:
        return f"R{v/1e9:,.2f}bn"
    if abs(v) >= 1e6:
        return f"R{v/1e6:,.1f}m"
    return f"R{v:,.0f}"


def style_doc(doc):
    n = doc.styles["Normal"]
    n.font.name = "Arial"
    n.font.size = Pt(11)
    for sid, size, color in [("Heading 1", 22, INK), ("Heading 2", 15, RED)]:
        s = doc.styles[sid]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color


def para(doc, text, size=11, color=None, bold=False, italic=False, align=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def pic(doc, path, caption, width=6.3):
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, caption, size=9, color=GREY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)


def page_break(doc):
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


def main():
    cat = pd.read_csv(os.path.join(DATA, "leatt_intelligent_category.csv"))
    chan = pd.read_csv(os.path.join(DATA, "leatt_intelligent_channel.csv"))
    mon = pd.read_csv(os.path.join(DATA, "leatt_intelligent_monthly.csv")).sort_values("month")
    prov = pd.read_csv(os.path.join(DATA, "leatt_intelligent_province.csv"))
    ab = pd.read_csv(os.path.join(DATA, "leatt_ab_test_results.csv"))
    comp = pd.read_csv(os.path.join(DATA, "leatt_competitor_analysis.csv"))
    exceptions = pd.read_csv(os.path.join(DATA, "leatt_audit_exceptions.csv"))
    controls = pd.read_csv(os.path.join(DATA, "leatt_data_controls.csv"))
    root_cause = pd.read_csv(os.path.join(DATA, "leatt_root_cause_playbook.csv"))
    signals = pd.read_csv(os.path.join(DATA, "leatt_decision_signal_rules.csv"))
    actions = pd.read_csv(os.path.join(DATA, "leatt_next_best_actions.csv"))
    initiatives = pd.read_csv(os.path.join(DATA, "leatt_prioritized_business_initiatives.csv"))
    risk_model = pd.read_csv(os.path.join(DATA, "leatt_ml_return_risk_metrics.csv")).iloc[0]
    risk_sample = pd.read_csv(os.path.join(DATA, "leatt_ml_return_risk_scores_sample.csv"))
    forecast = pd.read_csv(os.path.join(DATA, "leatt_revenue_forecast.csv"))

    revenue = cat["net_revenue_zar"].sum()
    margin = cat["gross_margin_zar"].sum()
    returns = cat["return_amount_zar"].sum()
    orders = int(chan["quantity"].sum())
    margin_pct = margin / revenue
    top_cat = cat.sort_values("net_revenue_zar", ascending=False).iloc[0]
    top_chan = chan.sort_values("net_revenue_zar", ascending=False).iloc[0]
    top_margin_cat = cat.sort_values("margin_rate", ascending=False).iloc[0]
    sig_wins = int((ab["Statistically significant"] == "Yes").sum())
    incremental = ab.loc[ab["Statistically significant"] == "Yes", "Estimated incremental revenue"].sum()
    flagged = mon[mon["anomaly_flag"] == True]
    flagged_margin = flagged["margin_rate"].mean()
    other_margin_pct = mon.loc[~mon.index.isin(flagged.index), "margin_rate"].mean()
    forecast_low = forecast["forecast_net_revenue_zar"].min()
    forecast_next = forecast.iloc[0]["forecast_net_revenue_zar"]

    doc = Document()
    style_doc(doc)

    # ---------- cover ----------
    for _ in range(5):
        doc.add_paragraph()
    para(doc, "LEATT GROWTH OS", size=36, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "A Fabric-powered ecommerce intelligence data story", size=15, color=RED,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    para(doc, "Microsoft Fabric · Data Factory · OneLake · Power BI · Python ML",
         size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "Anthony Apollis — July 2026", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Synthetic transaction data modelled on Leatt's public product catalog "
              "(leatt.com). Not affiliated with Leatt Corporation.",
         size=8, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    # ---------- 1. story ----------
    doc.add_heading("1. Is this profitable growth?", level=1)
    para(doc,
         f"Leatt sells motocross and adventure safety gear — helmets, body protection, boots, "
         f"gloves, goggles. Across {orders:,} modelled units sold, the catalog generates "
         f"{money(revenue)} in net revenue at a {margin_pct:.1%} gross margin, with "
         f"{returns/revenue:.1%} lost to returns. {top_cat['category']} is the largest category "
         f"({money(top_cat['net_revenue_zar'])}); {top_chan['channel']} the leading acquisition "
         f"channel ({money(top_chan['net_revenue_zar'])}).")
    pic(doc, os.path.join(IMG, "leatt_about_source.png"), "Figure 1 — Real Leatt storefront (leatt.com), the narrative anchor for this project.")
    pic(doc, os.path.join(IMG, "leatt_moto_boots_collection.png"), "Figure 2 — Real product catalog evidence: moto boots collection.", width=5.5)

    # ---------- 2. engine ----------
    doc.add_heading("2. The engine: what's real, and how data actually moves", level=1)
    para(doc,
         "A capacity, workspace, Lakehouse and Power BI report were created through the actual "
         "Fabric REST APIs — genuine, live items in the tenant, not a diagram of them:")
    for line in [
        "Workspace “Apollis” — e515bafe-7290-4832-ae1d-514be43a9d87",
        "Lakehouse “Leatt_BI_ML_Lakehouse” — dca60749-eaef-410e-9121-ea16eedbc975",
        "Power BI report “Leatt BI ML Executive Report” — 8c6988a7-fe5e-4fbe-abe9-ce7edd7fd63e",
        "Capacity “leattfabricf2” (SKU F2, South Africa North) — paused when not in active use",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    para(doc,
         "How the data actually moved: a Python script (src/fabric/upload_to_onelake.py) "
         "authenticated with an Azure CLI token and called OneLake's storage REST API directly "
         "(create-directory, append, flush) to push files into Bronze/Silver — this is the real, "
         "verified transfer mechanism, confirmed by matching byte counts on both ends: the "
         "90.7MB, 2,000,000-row transaction parquet; an 8.9MB, 11,354-variant product catalog; a "
         "13.5MB, 180,000-row synthetic customer file; and a 31.0MB customer ML-score file.")
    para(doc,
         "A Fabric Data Factory pipeline item (pl_leatt_million_row_lakehouse_load, "
         "9e82c185-e0ac-485d-9810-4bccdcfe6cf9) is genuinely registered in the workspace via the "
         "Fabric API, proving the integration exists — but it was not configured with copy "
         "activities and run; the physical move above happened via the script. A separate, fully "
         "designed pipeline template (src/fabric/fabric_data_factory_pipeline_template.json) "
         "specifies what a production build would do: a Copy activity pulling Leatt's live "
         "Shopify product API into Bronze, a Copy activity landing the transaction parquet into "
         "Bronze, and a Notebook activity building the Silver Delta table — presented here as a "
         "design, not as something proven to have executed.",
         size=9.5, italic=True, color=GREY)
    pic(doc, os.path.join(IMG, "fabric_data_factory_pipeline_blueprint.png"),
        "Figure 3 — The intended Bronze -> Silver -> Gold medallion flow (design target; Bronze/Silver landing is real, Gold is modelled locally today).")
    pic(doc, os.path.join(IMG, "leatt_star_schema_erd.png"),
        "Figure 4 — Star schema. Implemented today as a local SQLite warehouse; fabric_lakehouse_warehouse_ddl.sql is the equivalent starter DDL for Fabric.")
    para(doc,
         "Honest caveat: the Power BI report was created and bound to the semantic model via "
         "the real Fabric API (operation succeeded, 100% complete), but an automated service "
         "export only produced loading/sign-in screens rather than rendered visuals. Those blank "
         "exports are deliberately not presented here as visual proof — the report is real and "
         "reachable at its live URL; a rendered screenshot needs a signed-in interactive session.",
         size=9.5, italic=True, color=GREY)

    # ---------- 3. money map ----------
    doc.add_heading("3. The money map", level=1)
    para(doc,
         f"Growth quality, not just growth. {top_cat['category']} concentrates revenue but "
         f"{top_margin_cat['category']} carries the strongest margin "
         f"({top_margin_cat['margin_rate']:.1%}). November and December show consistent margin "
         f"compression in both years — {flagged_margin:.1%} average versus {other_margin_pct:.1%} "
         f"the rest of the year, a {other_margin_pct - flagged_margin:.1%} gap — a real seasonal "
         f"promotional pattern flagged by the anomaly-detection score in the gold layer, not noise.")
    pic(doc, os.path.join(CHARTS, "01_category.png"), "Figure 5 — Revenue and margin by category.")
    pic(doc, os.path.join(CHARTS, "02_monthly.png"), "Figure 6 — Monthly revenue and margin, with anomaly-flagged Nov/Dec months marked.")
    pic(doc, os.path.join(CHARTS, "03_channel.png"), "Figure 7 — Revenue by acquisition channel.")
    top_prov = prov.sort_values("net_revenue_zar", ascending=False).iloc[0]
    para(doc,
         f"Geographically, {top_prov['province']} leads at {money(top_prov['net_revenue_zar'])} "
         f"({top_prov['market_role'].lower()}); the full provincial breakdown follows.")
    pic(doc, os.path.join(CHARTS, "06_province.png"), "Figure 8 — Revenue by province.")

    # ---------- 4. issues, root causes and prevention ----------
    doc.add_heading("4. Issues, root causes and prevention — how this platform catches problems before they compound", level=1)
    para(doc,
         "A KPI dashboard that only reports numbers is half a BI platform. The other half is a "
         "structured way to ask \"why\", decide what to do about it, and catch the next one "
         "before it costs money. This chapter walks that loop end to end using the real signals "
         "generated from the gold layer.")

    para(doc, "4.1 Root-cause diagnostic playbook", bold=True, size=13, color=RED, space_after=4)
    para(doc, "When a headline number moves, this is the decision tree used to find out why "
              "before reacting:")
    for problem, grp in root_cause.groupby("problem"):
        para(doc, problem, bold=True, size=10.5, space_after=2)
        for _, r in grp.iterrows():
            para(doc, f"{r['diagnostic_branch']}: {r['question_to_ask']} → {r['likely_interpretation']}",
                 size=9.5, color=GREY, space_after=4)

    para(doc, "4.2 Decision signals monitored", bold=True, size=13, color=RED, space_after=4)
    para(doc, "Each signal below has a threshold rule; when a reading crosses it, an owner is "
              "already assigned to act — this is what makes the dashboard operational rather "
              "than decorative.")
    for _, r in signals.iterrows():
        para(doc, f"{r['signal']}: {r['current_reading']} — {r['health']}. {r['recommended_decision']} "
                  f"(owner: {r['owner']})", size=10, space_after=6)

    para(doc, "4.3 Prevention layer 1 — ML return-risk scoring", bold=True, size=13, color=RED, space_after=4)
    high_risk_rev = risk_sample["net_revenue_zar"].sum()
    high_risk_lines = len(risk_sample)
    top_risk_cat = risk_sample.groupby("category")["net_revenue_zar"].sum().sort_values(ascending=False).index[0]
    para(doc,
         f"A logistic regression model (AUC {risk_model['auc']:.2f}, recall {risk_model['recall']:.1%}) "
         f"scores every transaction line for return risk. The high-risk watchlist sample flags "
         f"{high_risk_lines:,} transaction lines worth {money(high_risk_rev)} at risk, concentrated "
         f"in {top_risk_cat} — giving CX and merchandising a ranked list to act on before the "
         f"return happens, rather than reading about it in next month's return-rate KPI.")
    pic(doc, os.path.join(CHARTS, "08_return_risk.png"), "Figure 9 — ML-flagged high-return-risk transactions by category.")

    para(doc, "4.4 Prevention layer 2 — forecast and anomaly detection", bold=True, size=13, color=RED, space_after=4)
    para(doc,
         f"A 6-month forward revenue forecast (next month: {money(forecast_next)}, trough "
         f"{money(forecast_low)}) combined with a monthly anomaly score lets the business see the "
         f"Nov/Dec margin dip coming rather than explaining it after the fact — the same 4 months "
         f"flagged by the anomaly score are exactly the 4 months with real margin compression.")
    pic(doc, os.path.join(CHARTS, "07_forecast.png"), "Figure 10 — Revenue trend, anomaly detection and 6-month forecast.")

    para(doc, "4.5 Next best actions and prioritized initiatives", bold=True, size=13, color=RED, space_after=4)
    for _, r in actions.iterrows():
        para(doc, f"{r['owner']}: {r['action']} — {r['rationale']}", size=10, space_after=6)
    para(doc, "Ranked by evidence-based priority score (value × confidence, effort-adjusted):",
         size=10, italic=True, color=GREY, space_after=4)
    for _, r in initiatives.sort_values("priority_score", ascending=False).iterrows():
        para(doc, f"[{r['priority_score']}] {r['initiative']} — est. value {r['estimated_value']}, "
                  f"{r['confidence']} confidence, {r['effort']} effort. Next step: {r['next_step']}",
             size=9.5, color=GREY, space_after=5)

    para(doc, "4.6 Prevention layer 3 — data controls (stopping bad data before it reaches the board)",
         bold=True, size=13, color=RED, space_after=4)
    for _, r in controls.iterrows():
        para(doc, f"{r['control_area']} — {r['control_name']}: {r['control_description']} "
                  f"({r['frequency']}, owner: {r['owner']})", size=9.5, color=GREY, space_after=4)

    # ---------- 5. growth loop ----------
    doc.add_heading("5. The growth loop: marketing, SEO, experimentation", level=1)
    para(doc,
         f"{sig_wins} of {len(ab)} A/B tests reached statistical significance, worth an "
         f"estimated {money(incremental)} in incremental revenue if rolled out.")
    pic(doc, os.path.join(CHARTS, "04_roas.png"), "Figure 11 — Marketing ROAS by channel.")
    pic(doc, os.path.join(CHARTS, "05_ab_tests.png"), "Figure 12 — A/B test results; red bars reached significance.")
    para(doc, "Competitive positioning:", bold=True, space_after=4)
    for _, r in comp.iterrows():
        para(doc, f"{r['Competitor']} — {r['Positioning']}. {r['Leatt opportunity']}", size=10, space_after=6)

    # ---------- 6. finance ----------
    doc.add_heading("6. Finance and governance: SAP-ready reconciliation", level=1)
    para(doc,
         "Ecommerce revenue reconciles to VAT, refunds and expected cash — the layer that lets "
         "Finance trust the BI numbers enough to post them against SAP Business One or SAP BW.")
    open_exc = (exceptions["status"] == "Open").sum()
    para(doc, f"{open_exc} of {len(exceptions)} sampled audit exceptions remain open, owned by "
              f"Finance Ops / Data Steward for resolution — each traceable back to the data "
              f"controls in section 4.6 that were designed to catch it.")
    pic(doc, os.path.join(IMG, "azure_portal_home_costs.png"), "Figure 13 — Real Azure portal session, cost dashboard visible.", width=5.8)

    # ---------- 7. cost control ----------
    doc.add_heading("7. Azure and Fabric: cost discipline", level=1)
    para(doc,
         "Fabric F2 capacity is paused by default and only resumed for active work, then "
         "suspended again immediately — verified via both the Azure CLI and the portal UI at "
         "every check this project has done.")
    para(doc, "Tags on the resource confirm intent: cost-control=delete-or-pause-after-upload, "
              "project=leatt-bi-ml, purpose=portfolio-proof. Full timestamped teardown log: "
              "docs/AZURE_COST_SHUTDOWN_PROOF.md.")
    para(doc, "Full source, data and reproduction steps: "
              "github.com/anthonyapollis/leatt-fabric-bi-ml-proof", size=9, color=GREY, italic=True)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    fix_zoom(OUT)
    print("written:", os.path.abspath(OUT))


def fix_zoom(path):
    """python-docx's default template writes <w:zoom> without the required
    w:percent attribute; patch it so strict OOXML validators pass."""
    import shutil
    import zipfile

    tmp = path + ".tmp"
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "word/settings.xml":
                text = data.decode("utf-8")
                if "<w:zoom" in text and "w:percent" not in text:
                    text = text.replace("<w:zoom ", '<w:zoom w:percent="100" ', 1)
                    text = text.replace("<w:zoom/>", '<w:zoom w:percent="100"/>', 1)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, path)


if __name__ == "__main__":
    main()
